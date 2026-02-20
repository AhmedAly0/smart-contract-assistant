"""
Document loading, extraction, and chunking utilities.
Pattern: Notebook 05 (ArxivLoader, UnstructuredFileLoader, RecursiveCharacterTextSplitter)
Pattern: Notebook 07 (metadata chunks, stub filtering, pre-References cutting)
"""

import os
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

from config import (
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    CHUNK_SEPARATORS,
    MIN_CHUNK_LENGTH,
    UPLOAD_DIR,
)


# ─── File Loaders ────────────────────────────────────────────────────────────

def load_pdf(filepath: str) -> str:
    """Extract text from a PDF file using PyMuPDF.
    Handles multi-page documents and strips excessive whitespace.
    """
    doc = fitz.open(filepath)
    text = ""
    for page in doc:
        text += page.get_text() + "\n"
    doc.close()
    return text.strip()


def load_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    doc = DocxDocument(filepath)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return text.strip()


def load_document(filepath: str) -> str:
    """Dispatcher: load text from a file based on its extension.

    Supports: .pdf, .docx, .txt
    Raises ValueError for unsupported formats.
    """
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return load_pdf(filepath)
    elif ext == ".docx":
        return load_docx(filepath)
    elif ext == ".txt":
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read().strip()
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use .pdf, .docx, or .txt")


# ─── Chunking ────────────────────────────────────────────────────────────────

def chunk_text(
    text: str,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    source: str = "uploaded_document",
) -> list[Document]:
    """Split text into chunks using RecursiveCharacterTextSplitter.
    Pattern: Notebook 05/07 — chunking with recommended separators and stub filtering.

    Args:
        text: Raw document text.
        chunk_size: Max characters per chunk (default from config).
        chunk_overlap: Overlap between chunks (default from config).
        source: Source name for metadata.

    Returns:
        List of LangChain Document objects with metadata.
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=CHUNK_SEPARATORS,
    )

    # Create a single Document, then split
    doc = Document(page_content=text, metadata={"source": source, "Title": source})
    chunks = splitter.split_documents([doc])

    # Filter out stubs — chunks too small to be useful (Notebook 07 pattern)
    chunks = [c for c in chunks if len(c.page_content) >= MIN_CHUNK_LENGTH]

    return chunks


def prepare_metadata_chunks(doc_name: str, chunks: list[Document]) -> list[Document]:
    """Add metadata and summary chunks to the document set.
    Pattern: Notebook 07 — prepend metadata chunks for better retrieval context.

    Adds:
    1. A document overview chunk with name and total chunk count.
    2. Updates all chunk metadata with consistent source info.
    """
    # Overview chunk
    overview = Document(
        page_content=(
            f"This document is titled '{doc_name}'. "
            f"It has been split into {len(chunks)} chunks for analysis. "
            f"Ask questions about this document to get answers with source citations."
        ),
        metadata={"source": doc_name, "Title": doc_name, "type": "overview"},
    )

    # Ensure all chunks have consistent metadata
    for i, chunk in enumerate(chunks):
        chunk.metadata["source"] = doc_name
        chunk.metadata["Title"] = doc_name
        chunk.metadata["chunk_index"] = i

    return [overview] + chunks


# ─── Full Ingestion Pipeline ─────────────────────────────────────────────────

def ingest_document(filepath: str, doc_name: Optional[str] = None) -> list[Document]:
    """Full ingestion pipeline: load → chunk → add metadata.

    Args:
        filepath: Path to the document file.
        doc_name: Optional display name. Defaults to filename.

    Returns:
        List of LangChain Documents ready for embedding and vector store insertion.
    """
    if doc_name is None:
        doc_name = os.path.basename(filepath)

    # Load
    text = load_document(filepath)

    if not text.strip():
        raise ValueError(f"Document '{doc_name}' is empty or could not be parsed.")

    # Chunk
    chunks = chunk_text(text, source=doc_name)

    if not chunks:
        raise ValueError(
            f"Document '{doc_name}' produced no valid chunks. "
            f"It may be too short (min chunk length: {MIN_CHUNK_LENGTH} chars)."
        )

    # Add metadata
    chunks = prepare_metadata_chunks(doc_name, chunks)

    print(f"[Ingestion] '{doc_name}': {len(chunks)} chunks created "
          f"(chunk_size={CHUNK_SIZE}, overlap={CHUNK_OVERLAP})")

    return chunks


def save_uploaded_file(file_bytes: bytes, filename: str) -> str:
    """Save uploaded file bytes to the data directory. Returns the saved filepath."""
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    filepath = os.path.join(UPLOAD_DIR, filename)
    with open(filepath, "wb") as f:
        f.write(file_bytes)
    return filepath
